"""POST /v1/knowledge/import-file 端点的集成测试。

复用 test_api.py 的 client fixture（sqlite 临时库、向量索引关掉）。
"""
from __future__ import annotations

from io import BytesIO

import pytest
from fastapi.testclient import TestClient


@pytest.fixture
def client(tmp_path, monkeypatch):
    monkeypatch.setenv("KB_BACKEND", "sqlite")
    monkeypatch.setenv("SQLITE_PATH", str(tmp_path / "test.db"))
    monkeypatch.setenv("VECTOR_ENABLED", "0")

    cfg_path = tmp_path / "config.toml"
    cfg_path.write_text("[server]\nport = 18000\n", encoding="utf-8")
    monkeypatch.setenv("KB_CONFIG_TOML_PATH", str(cfg_path))

    from app.main import _repo_singleton_sqlite, _repo_singleton_postgres
    _repo_singleton_sqlite.cache_clear()
    _repo_singleton_postgres.cache_clear()

    from app.main import app
    return TestClient(app)


def _post_file(client: TestClient, filename: str, content: bytes, **form: str):
    data = {
        "project": "proj-import",
        "domain": "work",
        **form,
    }
    files = {"file": (filename, BytesIO(content), "application/octet-stream")}
    return client.post("/v1/knowledge/import-file", data=data, files=files)


def test_markdown_upload_succeeds(client: TestClient) -> None:
    md = b"# Hello Title\n\nThis is body content for the document.\n"
    resp = _post_file(client, "demo.md", md)
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["knowledge_item_id"]
    assert body["version"] == 1


def test_title_inferred_from_first_h1(client: TestClient) -> None:
    md = b"# Auto Inferred\n\nBody."
    resp = _post_file(client, "anything.md", md)
    assert resp.status_code == 200
    item_id = resp.json()["knowledge_item_id"]

    detail = client.get(f"/v1/knowledge/items/{item_id}")
    assert detail.status_code == 200
    assert detail.json()["title"] == "Auto Inferred"


def test_explicit_title_overrides_inference(client: TestClient) -> None:
    md = b"# Inferred From Body\n\nContent."
    resp = _post_file(client, "x.md", md, title="Explicit Title")
    assert resp.status_code == 200
    item_id = resp.json()["knowledge_item_id"]

    detail = client.get(f"/v1/knowledge/items/{item_id}")
    assert detail.json()["title"] == "Explicit Title"


def test_unsupported_extension_returns_415(client: TestClient) -> None:
    resp = _post_file(client, "logo.png", b"\x89PNG\r\n\x1a\n")
    assert resp.status_code == 415
    assert "不支持" in resp.json()["detail"] or "support" in resp.json()["detail"].lower()


def test_empty_document_returns_400(client: TestClient) -> None:
    resp = _post_file(client, "empty.md", b"   \n\n  ")
    assert resp.status_code == 400


def test_missing_project_returns_422(client: TestClient) -> None:
    files = {"file": ("a.md", BytesIO(b"# x\nbody"), "application/octet-stream")}
    resp = client.post(
        "/v1/knowledge/import-file",
        data={"domain": "work"},
        files=files,
    )
    assert resp.status_code == 422


def test_invalid_domain_returns_400(client: TestClient) -> None:
    resp = _post_file(client, "a.md", b"# x\nbody", domain="invalid-domain")
    assert resp.status_code == 400


def test_docx_upload_parses_paragraphs(client: TestClient, tmp_path) -> None:
    """构造一个真实 .docx 验证 python-docx 解析链。"""
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_paragraph("第一段：这是文档第一段内容。")
    doc.add_paragraph("第二段：这是第二段。")
    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))

    resp = _post_file(client, "sample.docx", docx_path.read_bytes(), title="My DOCX")
    assert resp.status_code == 200, resp.text
    item_id = resp.json()["knowledge_item_id"]

    detail = client.get(f"/v1/knowledge/items/{item_id}")
    assert detail.status_code == 200
    body = detail.json()
    assert body["title"] == "My DOCX"
    assert "第一段" in body["content_markdown"]
    assert "第二段" in body["content_markdown"]
